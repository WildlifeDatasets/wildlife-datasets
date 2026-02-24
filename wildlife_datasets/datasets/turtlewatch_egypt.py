import os
import re
from collections.abc import Callable, Sequence
from typing import cast

import numpy as np
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph
from tqdm import tqdm

from .datasets import WildlifeDataset, utils
from .general import Dataset_Metadata
from .utils import load_segmentation as utils_load_segmentation
from .utils import strip_suffixes

identity_replace = {
    "": np.nan,
    "l": "unknown",
    "r": "unknown",
    "42016 ni ei ni": "4-2016 ni ei ni",
    "72016 ni ei baby": "7-2016 ni ei baby",
    "112016 ni baby barn r": "11-2016 ni baby barn",
    "1noid": "unknown",
    "2noid": "unknown",
    "_amal kabir": "amal kabir",
    "anette": "annette",
    "annvd": "unknown",
    "arf": "unknown",
    "asunk1": "mara",
    "asunk2": "enea",
    "asunk5": "fame",
    "asukn8": "asunk8",
    "barney": "rileside",
    "carn ei1": "carn1 ei",
    "dabb1 ei": "cassiopea",
    "dabb3": "dabb3 ei",
    "dabb3ei": "dabb3 ei",
    "dabb6ei": "momo",
    "dabb6 ei": "momo",
    "elph2": "elph2 ei",
    "elph3 ei": "jamal",
    "egunk3": "zakaria",
    "egunk8": "tsunami",
    "egunk10": "alice",
    "end june gerryandwout van koot-r 2": "unknown",
    "female no id": "unknown",
    "franncina": "francina",
    "gt-0006 karen": "karen",
    "halhala1": "halhala 1",
    "her1 ei": "boma",
    "her2": "her2 ei",
    "her2ei": "her2 ei",
    "her3": "her4 ei",
    "her3ei": "her4 ei",
    "her3 ei": "her4 ei",
    "her4": "her4 ei",
    "her4ei": "her4 ei",
    "her5ei": "her5 ei",
    "her6ei": "her6 ei",
    "her8": "her8 ei",
    "her9 ei": "luna",
    "her hunk1": "hunk1",
    "her hunk2": "ansofie",
    "her hunk6": "hunk6",
    "her hunk10": "layali",
    "hunk11": "niko",
    "hunk14": "chris",
    "hunk15": "pearl",
    "hunk16": "norman",
    "hunk17": "christopher columbus",
    "img": "unknown",
    "inbound1434235978309711639": "unknown",
    "inbound2545461428103863063": "unknown",
    "inbound2900475970750989766": "unknown",
    "inbound4209230724466990989": "unknown",
    "injured": "unknown",
    "juve maxim": "maxim",
    "kriya": "krilya",
    "ladilsav": "ladislav",
    "maas": "maas ei 1",
    "maas el1": "maas ei 1",
    "maasei 1": "maas ei 1",
    "magherita": "margherita",
    "no": "unknown",
    "nosureidher3ei": "unknown",
    "petra-3": "petra",
    "perolaf": "per-olaf",
    "per olaf": "per-olaf",
    "red sea scuba diving - green turtle": "unknown",
    "rhan": "rhian",
    "saskia": "saskia ch",
    "sata1ei": "sata1 ei",
    "shab 1": "shab1 ei",
    "tondoba": "louisa",
    "tond1ei": "tond1 ei",
    "tond2ei": "tond2 ei",
    "toro": "toro1 ei",
    "toro1": "toro1 ei",
    "unk1": "phealyn",
    "unk2": "sebastian",
    "unk3": "didier",
    "unk5": "petra",
    "unk7": "jan",
    "unk8": "sascha",
    "unk10": "lin",
    "unk14": "maxim",
    "unk15": "kariman",
    "unk16": "andrea",
    "unk19": "altea",
    "unk20": "platon",
    "unk21": "ruud",
    "unk24": "karim",
    "unk27": "honey",
    "unk30": "carmen",
    "unk31": "lisa",
    "unk34": "manu",
    "unkei": "unknown",
    "wdla1": "wdla1 ei",
}


def get_encounter_name(x: str) -> str | None:
    i = x.lower().rfind("sighting")
    if i == -1:
        return None
    i = x.find(os.path.sep, i)
    if i == -1:
        return x
    else:
        return x[:i]


def merge_codes(df: pd.DataFrame, index, individuals: list[str]) -> list[str]:
    xs = []
    for name in df["file"]:
        x = code_to_info(name, individuals)[index]
        xs.append(x)
    return list(pd.Series(xs).dropna().unique())


def get_code(xs: Sequence[str], name: str = "variables") -> str | None:
    xs = list(xs)
    for y in ["", "unknown"]:
        if len(xs) > 1 and y in xs:
            xs.remove(y)
    if len(xs) >= 1:
        return xs[0]
    else:
        return None


def fix_identity(x: str | None, individuals: list[str]) -> str | None:
    if pd.isnull(x):
        return x
    x = x.strip().lower()
    if x.startswith("no id") or x.startswith("noid"):
        return "unknown"

    ys = [
        x,
        " ".join(x.split(" ")[:2]),
        " ".join(x.split(" ")[:-1]),
        x.split(" ")[0],
        x.split("-")[0],
        "-".join(x.split("-")[:-1]),
        ".".join(x.split(".")[:-1]),
        x.split("_")[0],
    ]
    for y in ys:
        if y in individuals:
            return y
        for str1, str2 in identity_replace.items():
            if y == str1 and (str2 == "unknown" or str2 in individuals):
                return str2

    return x


def is_int(x) -> bool:
    try:
        return int(x) == float(x)
    except (TypeError, ValueError):
        return False


def code_to_info(x: str, individuals: list[str]) -> tuple[str | None, ...]:
    codes_morning = ["am", "am1"]
    codes_afternoon = ["pm"]
    codes_time = codes_morning + codes_afternoon

    x = os.path.splitext(x)[0]
    x = x.lower()
    x_split = x.split("_")
    if len(x_split) >= 10:
        identity = x_split[0].strip()
        orientation = x_split[1]
        leader = x_split[2]

        # Extract date
        date_year = x_split[3]
        date_month = x_split[4]
        date_day = x_split[5]
        date = None
        if is_int(date_year) and is_int(date_month) and is_int(date_day):
            date_year = int(date_year)
            date_month = int(date_month)
            date_day = int(date_day)
            if int(date_year) >= 2000 and (1 <= int(date_month) <= 12) and (1 <= int(date_day) <= 31):
                date = f"{date_year:04d}-{date_month:02d}-{date_day:02d}"

        # Extract place
        if x_split[6].strip() in codes_time:
            i_place, i_noon, i_hour = 8, 6, 7
        elif x_split[8].strip() in codes_time:
            i_place, i_noon, i_hour = 6, 8, 7
        else:
            i_place, i_noon, i_hour = 6, 7, 8
        place = x_split[i_place]

        # Extract hour
        hour = None
        if x_split[i_noon].strip() in codes_morning:
            hour_candidate = x_split[i_hour].split(".")[0]
            if is_int(hour_candidate):
                hour = np.mod(int(hour_candidate), 12)
        elif x_split[i_noon].strip() in codes_afternoon:
            hour_candidate = x_split[i_hour]
            if is_int(hour_candidate):
                hour = 12 + np.mod(int(hour_candidate), 12)

        # Extract author
        author = x_split[9]
        for x in ["-", "(", "."]:
            author = author.split(x)[0]
        author = re.sub(r"\d+$", "", author).strip()
    elif len(x_split) >= 5:
        identity = x_split[0].strip()
        orientation, leader, date, place, hour, author = None, None, None, None, None, None
    else:
        identity, orientation, leader, date, place, hour, author = (
            None,
            None,
            None,
            None,
            None,
            None,
            None,
        )
    identity = fix_identity(identity, individuals)
    return identity, orientation, leader, date, place, hour, author


def info_to_code(
    identity: str, orientation: str, leader: str, date: str, place: str, hour: int | None = None, author: str = ""
) -> str:
    date_split = str(date).split("-")
    year = date_split[0]
    month = date_split[1]
    day = date_split[2]
    if pd.isnull(hour):
        hour1 = ""
        hour2 = ""
    else:
        if hour < 12:
            hour1 = "AM"
            hour2 = int(hour)
        else:
            hour1 = "PM"
            hour2 = int(hour) - 12
    code1 = f"{identity}_{orientation}_{leader}_{year}_{month}_{day}_{place}_{hour1}_{hour2}".upper()
    code2 = author.title()
    return f"{code1}_{code2}"


class TurtlewatchEgypt_Base(WildlifeDataset):
    @classmethod
    def _download(cls, **kwargs) -> None:
        pass

    @classmethod
    def _extract(cls, **kwargs) -> None:
        pass

    def extract_info(self, i: int) -> tuple[str | None, ...]:
        path = self.df.at[i, "path"]
        assert isinstance(path, str)
        return code_to_info(path.split(os.path.sep)[-1], self.individuals)

    def extract_code(self, i: int) -> str:
        df_row = self.df.iloc[i]
        identity = df_row["identity"]
        orientation = df_row["label"]
        leader = df_row["leader"]
        date = df_row["date"]
        place = df_row["place"] if not pd.isnull(df_row["place"]) else ""
        hour = df_row["hour"]
        author = df_row["author"] if not pd.isnull(df_row["author"]) else ""
        return info_to_code(identity, orientation, leader, date, place, hour=hour, author=author)

    def load_individuals(self, file_name: str | None = None) -> None:
        if file_name is None:
            file_name = f"{os.path.dirname(os.path.abspath(__file__))}/individuals.csv"
        if not os.path.exists(file_name):
            raise ValueError(f"File does not exist: {file_name}")
        individuals = pd.read_csv(file_name)
        individuals = individuals["Common_name"].to_numpy()
        individuals = [x.lower().strip() for x in individuals]
        self.individuals = [strip_suffixes(x, [" C", " (DEAD)"]) for x in individuals]


class TurtlewatchEgypt_Master(TurtlewatchEgypt_Base):
    def create_catalogue(self, file_name: str | None = None) -> pd.DataFrame:
        assert self.root is not None
        self.load_individuals(file_name=file_name)
        data = utils.find_images(self.root)

        # Get identity
        data["identity"] = data["file"].apply(lambda x: fix_identity(x.lower(), self.individuals))

        # Get orientation
        data["date"] = data["file"].apply(lambda x: code_to_info(os.path.basename(x), self.individuals)[3])
        orientation = data["file"].apply(lambda x: code_to_info(os.path.basename(x), self.individuals)[1])
        idx = orientation.isnull()
        orientation[idx] = data.loc[idx, "file"].apply(lambda x: os.path.basename(x).split(".")[-2][-2:])
        orientation = orientation.apply(lambda x: x.strip().lower())
        orientation = orientation.apply(
            lambda x: x if x in ["l", "r", "b", "c", "t", "alf", "arf", "rlf", "rrf"] else None
        )
        data["orientation"] = orientation

        # Add path and corresponding image_id
        data["path"] = data["path"] + os.path.sep + data["file"]
        data["image_id"] = utils.get_persistent_id(data["path"])

        # Finalize the dataframe
        data = data.drop("file", axis=1)
        return self.finalize_catalogue(data)


class TurtlewatchEgypt_New(TurtlewatchEgypt_Base):
    def create_catalogue(self, load_segmentation: bool = False, file_name: str | None = None) -> pd.DataFrame:

        assert self.root is not None
        self.load_individuals(file_name=file_name)
        data = utils.find_images(self.root)

        # Ignoring data starting with '.'
        mask = data["file"].str.startswith(".")
        data = data[~mask]

        # Get full file names
        data["path_full"] = data["path"] + os.path.sep + data["file"]

        # Adding artificial encounters
        data["encounter_name"] = data["path_full"].apply(lambda x: get_encounter_name(x))
        idx = data["encounter_name"].isnull()
        if sum(idx) > 0:
            for folder, df_folder in data[idx].groupby("path"):
                assert isinstance(folder, str)
                data.loc[df_folder.index, "encounter_name"] = folder.lower()

        # Sort data
        data = data.sort_values("encounter_name")
        data = data.reset_index(drop=True)

        # Get encounter_id
        data["encounter_id"] = (data["encounter_name"] != data["encounter_name"].shift()).cumsum()
        assert data["encounter_id"].nunique() == data["encounter_name"].nunique()

        # Preallocate columns to be able to handle strings and nans without warnings
        data["identity"] = pd.Series([None] * len(data), dtype="object")
        data["leader"] = pd.Series([None] * len(data), dtype="object")
        data["place"] = pd.Series([None] * len(data), dtype="object")
        data["author"] = pd.Series([None] * len(data), dtype="object")
        data["date"] = pd.Series([None] * len(data), dtype="object")

        for _, df_encounter in data.groupby("encounter_id"):
            # Extract information from code
            identities = merge_codes(df_encounter, 0, self.individuals)
            leaders = merge_codes(df_encounter, 2, self.individuals)
            dates = merge_codes(df_encounter, 3, self.individuals)
            places = merge_codes(df_encounter, 4, self.individuals)
            hours = merge_codes(df_encounter, 5, self.individuals)
            authors = merge_codes(df_encounter, 6, self.individuals)
            # Add individuals if empty
            if len(identities) == 0:
                for name in df_encounter["path"]:
                    for name_split1 in name.split(os.path.sep):
                        name_split2 = name_split1.lower().split(" ")
                        if name_split2[0] == "sighting" and len(name_split2) >= 3:
                            identities.append(name_split2[2])
                identities = pd.Series(identities).dropna().unique().tolist()
            # Save codes
            data.loc[df_encounter.index, "identity"] = get_code(identities, name="identities")
            data.loc[df_encounter.index, "leader"] = get_code(leaders, name="leaders")
            data.loc[df_encounter.index, "place"] = get_code(places, name="places")
            data.loc[df_encounter.index, "hour"] = get_code(hours, name="hours")
            data.loc[df_encounter.index, "author"] = get_code(authors, name="authors")
            data.loc[df_encounter.index, "date"] = get_code(dates, name="dates")

        # Fix the column names
        data = data.reset_index(drop=True)
        data = data.drop(["path", "file", "encounter_name"], axis=1)
        data = data.rename({"path_full": "path"}, axis=1)

        # Fix unknown individuals
        data.loc[data["identity"].isnull(), "identity"] = "unknown"

        # Add persistent image_id
        data["image_id"] = utils.get_persistent_id(data["path"])

        # Load segmentation
        if load_segmentation:
            data = utils_load_segmentation(data, os.path.join(self.root, "segmentation.csv"))
        return self.finalize_catalogue(data)


############################################
# Functions related to downloads
############################################


class TurtlewatchEgypt_Citizen(Dataset_Metadata):
    @classmethod
    def _download(cls, data: pd.DataFrame | None = None, transform: Callable | None = None) -> None:
        # Transform the data into the required form
        assert data is not None
        data = load_citizen_data(data)
        if transform is not None:
            data = transform(data)
        assert isinstance(data, pd.DataFrame)

        # Go through the rows and download data
        metadata = pd.DataFrame()
        for encounter, (_, d) in tqdm(enumerate(data.iterrows()), total=len(data)):
            urls = d["upload-file-731"].split("\n")
            folder = d["folder"]
            sighting = d["sighting"]
            folder_full = os.path.join(folder, f"SIGHTING #{sighting}")

            save_paths = download_files(urls, folder_full)
            save_paths = [os.path.relpath(p, ".") for p in save_paths]

            create_info(d, folder_full)

            metadata_part = {
                "path": save_paths,
                "identity": "unknown",
                "encounter_id": encounter,
            }
            metadata = pd.concat((metadata, pd.DataFrame(metadata_part)))

        # Assign image_id
        metadata = metadata.reset_index(drop=True)
        metadata["image_id"] = utils.get_persistent_id(metadata["path"])
        metadata.to_csv("metadata.csv", index=False)

    @classmethod
    def _extract(cls, **kwargs) -> None:
        pass


IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp", ".gif", ".heic", ".heif", ".avif")


def load_citizen_data(data: pd.DataFrame) -> pd.DataFrame:
    # Convert dates
    data["submit_time"] = pd.to_datetime(data["submit_time"])
    data["date-227"] = pd.to_datetime(data["date-227"])

    # Merge multiple date options
    data["date"] = data["submit_time"].combine_first(data["date-227"])
    data["year"] = data["date"].dt.year
    data["month"] = data["date"].dt.month
    data["day"] = data["date"].dt.day

    # Merge multiple author and emails options
    data["author"] = data["NOME"].combine_first(data["your-name"])
    data["email"] = data["EMAIL"].combine_first(data["email-185"])

    # Fill nans
    data["author"] = data["author"].fillna("unknown")
    data["email"] = data["email"].fillna("unknown")

    # Get folder and sightings
    data["folder"] = [get_folder(d) for _, d in data.iterrows()]
    for _, data_folder in data.groupby(["folder"]):
        data.loc[data_folder.index, "sighting"] = list(range(1, len(data_folder) + 1))
    data["sighting"] = data["sighting"].astype(int)

    return data


def get_folder(d: pd.Series) -> str:
    year = d["year"]
    month = d["month"]
    day = d["day"]

    folder1 = f"{year}_{month:02d}_{day:02d}"
    folder2 = d["author"]
    return f"{folder1}/{folder2}"


def download_files(urls: list[str], download_folder: str, exts: tuple[str, ...] = IMAGE_EXTENSIONS) -> list[str]:
    os.makedirs(download_folder, exist_ok=True)

    save_paths = []
    for url in urls:
        if not url.lower().endswith(exts):
            print(f"Skipping non-image url: {url}")
            continue
        file_name = url.split("/")[-1]
        save_path = os.path.join(download_folder, file_name)
        save_paths.append(save_path)
        if os.path.exists(save_path):
            continue
        try:
            response = requests.get(url, timeout=20)
            if response.status_code == 200:
                with open(save_path, "wb") as f:
                    f.write(response.content)
            else:
                print(f"Failed: {url}")
        except Exception as e:
            print(f"Error downloading {url}: {e}")
    return save_paths


def add_run_break(p: Paragraph, text1: str, text2: str | None = None) -> None:
    if not pd.isnull(text2):
        r = p.add_run(f"{text1}: {text2}")
    else:
        r = p.add_run(f"{text1}")
    r.add_break()


def create_info(d: pd.Series, save_folder: str) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style = cast(ParagraphStyle, style)

    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    p = doc.add_paragraph()
    add_run_break(p, "REQUIRED DATA")
    add_run_break(p, "From", d["email"])
    add_run_break(p, "Photographer", d["author"])
    add_run_break(p, "Date", f"{d['year']}-{d['month']:02d}-{d['day']:02d}")
    add_run_break(p, "Town", d["town-785"])
    add_run_break(p, "Location", d["location-785"])
    add_run_break(p, "")
    add_run_break(p, "OPTIONAL DATA")
    add_run_break(p, "Dive centre/independent", d["dive-785"])
    add_run_break(p, "Time", d["time-299"])
    add_run_break(p, "Depth", d["depth-364"])
    add_run_break(p, "Temperature", d["degrees-364"])
    add_run_break(p, "Activity", d["radio-602"])
    add_run_break(p, "Species", d["species-954"])
    add_run_break(p, "Size", d["size-648"])
    add_run_break(p, "Sex", d["sex-786"])
    add_run_break(p, "Comments", d["textarea-268"])
    if d["acceptance-359"] == 1:
        add_run_break(p, "")
        add_run_break(p, "DATA TREATMENT")
        add_run_break(
            p,
            "I allow TurtleWatch Egypt 2.0 to use my digital contents (photos and videos) and the data entered in this form for didactic, educational and scientific use: Yes",
        )
        add_run_break(p, "")
        add_run_break(
            p,
            "I allow TurtleWatch Egypt 2.0 to use my digital contents (photos and videos) and the data entered in this form for marketing and advertising use (social media, magazines, ..): Yes",
        )
        add_run_break(p, "")
        add_run_break(p, "Accettato: I authorize the treatment and management of personal data.")

    doc.save(f"{save_folder}/info.docx")
